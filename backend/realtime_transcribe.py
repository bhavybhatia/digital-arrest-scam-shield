import os
import sys
import queue
import threading
import numpy as np
from docx import Document
from datetime import datetime

# Audio arrives as 16-bit mono PCM chunks pushed onto audio_queue by the
# WebSocket receiver in app.py (captured client-side in the browser) rather
# than read from a local microphone here. RATE/CHUNK_DURATION still define
# the expected chunk shape: app.py uses them to size the "full" chunk in
# bytes so it can tell a final, shorter chunk (call ended mid-recording)
# from a normal one.
RATE = 16000
CHUNK_DURATION = 10   # seconds per chunk
OUTPUT_PATH = "transcription.docx"

audio_queue = queue.Queue()
stop_event = threading.Event()
doc_lock = threading.Lock()

CHUNK_BYTES = RATE * 2 * CHUNK_DURATION  # 16-bit mono PCM


class ChunkAssembler:
    """Buffers raw PCM16 bytes streamed in over the network into fixed
    ~CHUNK_DURATION-second pieces, mirroring what the old microphone-based
    record_audio() used to hand to audio_queue directly."""

    def __init__(self):
        self._buffer = bytearray()

    def add(self, raw_bytes):
        self._buffer.extend(raw_bytes)
        chunks = []
        while len(self._buffer) >= CHUNK_BYTES:
            chunks.append(bytes(self._buffer[:CHUNK_BYTES]))
            del self._buffer[:CHUNK_BYTES]
        return chunks

    def flush(self):
        if not self._buffer:
            return None
        chunk = bytes(self._buffer)
        self._buffer = bytearray()
        return chunk


def init_doc():
    if os.path.exists(OUTPUT_PATH):
        print(f"Existing file found: {OUTPUT_PATH} — transcripts will be appended.")
    else:
        doc = Document()
        doc.add_heading("Live Call Transcription Log", level=1)
        doc.save(OUTPUT_PATH)
        print(f"Created new file: {OUTPUT_PATH}")


def _load_or_create_doc():
    try:
        return Document(OUTPUT_PATH)
    except Exception:
        # transcription.docx missing or corrupted (e.g. an earlier crash
        # truncated it mid-write) - recreate it rather than permanently
        # breaking every future append.
        doc = Document()
        doc.add_heading("Live Call Transcription Log", level=1)
        return doc


def append_to_doc(timestamp, text, note=None, risk_score=None, risk_label=None, risk_status=None):
    with doc_lock:
        doc = _load_or_create_doc()
        heading = f"{timestamp}  ({note})" if note else timestamp
        doc.add_paragraph(heading, style="Intense Quote")
        doc.add_paragraph(text)
        if risk_status is not None:
            risk_line = f"Risk: {risk_status}  ({risk_score}/100)"
            if risk_label:
                risk_line += f"  — intent: {risk_label}"
            doc.add_paragraph(risk_line, style="Intense Quote")
        doc.save(OUTPUT_PATH)


def mark_call_event(label):
    """Write a call start/end marker into the docx so the fixed ~10s
    transcript chunks are grouped per call instead of reading as one
    continuous, inconsistently-spaced log with unexplained gaps."""
    with doc_lock:
        doc = _load_or_create_doc()
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        doc.add_heading(f"— {label} — {timestamp} —", level=2)
        doc.save(OUTPUT_PATH)


def transcribe_audio(model, on_chunk=None):
    while not stop_event.is_set() or not audio_queue.empty():
        try:
            timestamp, raw, is_partial = audio_queue.get(timeout=1)
        except queue.Empty:
            continue

        try:
            audio_np = np.frombuffer(raw, dtype=np.int16).astype(np.float32) / 32768.0
            rms = np.sqrt(np.mean(audio_np ** 2))
            actual_duration = len(audio_np) / RATE
            note = (
                f"{actual_duration:.1f}s partial chunk — call ended mid-recording"
                if is_partial else None
            )

            if rms < 0.001:
                print(f"[{timestamp}] Chunk skipped (silence).")
                continue

            audio_np = audio_np / (np.max(np.abs(audio_np)) + 1e-9) * 0.9
            result = model.transcribe(audio_np, language="en", fp16=False)
            text = result["text"].strip()

            if text:
                print(f"[{timestamp}] >> {text}")
                sys.stdout.flush()

                risk_score = risk_label = risk_status = None
                if on_chunk:
                    chunk_result = on_chunk(timestamp, text)
                    if chunk_result:
                        risk_score, risk_label, risk_status = chunk_result

                append_to_doc(
                    timestamp, text, note=note,
                    risk_score=risk_score, risk_label=risk_label, risk_status=risk_status,
                )
                print(f"[{timestamp}] Appended to {OUTPUT_PATH}")
            else:
                print(f"[{timestamp}] No speech detected in chunk.")
        except Exception as exc:
            # A single bad chunk (transcription error, scoring error, a
            # print() encoding issue, ...) must not kill this thread - that
            # would silently stop all further transcription for the rest of
            # the call with no visible error to the user.
            print(f"[{timestamp}] Chunk processing failed: {exc}")
