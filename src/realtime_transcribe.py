import os
import sys
import queue
import threading
import numpy as np
import pyaudio
import whisper
from docx import Document
from datetime import datetime

CHUNK = 1024
FORMAT = pyaudio.paInt16
CHANNELS = 1
RATE = 16000
CHUNK_DURATION = 10   # seconds per chunk
OUTPUT_PATH = "transcription.docx"

audio_queue = queue.Queue()
stop_event = threading.Event()
doc_lock = threading.Lock()


def init_doc():
    if os.path.exists(OUTPUT_PATH):
        print(f"Existing file found: {OUTPUT_PATH} — transcripts will be appended.")
    else:
        doc = Document()
        doc.add_heading("Live Call Transcription Log", level=1)
        doc.save(OUTPUT_PATH)
        print(f"Created new file: {OUTPUT_PATH}")


def append_to_doc(timestamp, text):
    with doc_lock:
        doc = Document(OUTPUT_PATH)
        doc.add_paragraph(timestamp, style="Intense Quote")
        doc.add_paragraph(text)
        doc.save(OUTPUT_PATH)


def record_audio():
    p = pyaudio.PyAudio()
    stream = p.open(format=FORMAT, channels=CHANNELS, rate=RATE,
                    input=True, frames_per_buffer=CHUNK)
    frames_per_chunk = int(RATE / CHUNK * CHUNK_DURATION)

    while not stop_event.is_set():
        frames = []
        for _ in range(frames_per_chunk):
            if stop_event.is_set():
                break
            frames.append(stream.read(CHUNK, exception_on_overflow=False))
        if frames:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            audio_queue.put((timestamp, b''.join(frames)))

    stream.stop_stream()
    stream.close()
    p.terminate()


def transcribe_audio(model):
    while not stop_event.is_set() or not audio_queue.empty():
        try:
            timestamp, raw = audio_queue.get(timeout=1)
        except queue.Empty:
            continue

        audio_np = np.frombuffer(raw, dtype=np.int16).astype(np.float32) / 32768.0
        rms = np.sqrt(np.mean(audio_np ** 2))

        if rms < 0.001:
            print(f"[{timestamp}] Chunk skipped (silence).")
            continue

        audio_np = audio_np / (np.max(np.abs(audio_np)) + 1e-9) * 0.9
        result = model.transcribe(audio_np, language="en", fp16=False)
        text = result["text"].strip()

        if text:
            print(f"[{timestamp}] >> {text}")
            sys.stdout.flush()
            append_to_doc(timestamp, text)
            print(f"[{timestamp}] Appended to {OUTPUT_PATH}")
        else:
            print(f"[{timestamp}] No speech detected in chunk.")


if __name__ == "__main__":
    # Step 1: Check / create docx
    init_doc()

    # Load model
    print("\nLoading Whisper model (tiny)...")
    model = whisper.load_model("tiny")

    print("\nLive transcription started. Speak freely. Press Ctrl+C to stop.\n")

    recorder = threading.Thread(target=record_audio, daemon=True)
    transcriber = threading.Thread(target=transcribe_audio, args=(model,), daemon=True)

    recorder.start()
    transcriber.start()

    try:
        while recorder.is_alive():
            recorder.join(timeout=0.5)
    except KeyboardInterrupt:
        print("\nStopping... finishing pending transcriptions.")
        stop_event.set()

    recorder.join()
    transcriber.join()
    print(f"\nDone. All transcripts saved to {OUTPUT_PATH}")
